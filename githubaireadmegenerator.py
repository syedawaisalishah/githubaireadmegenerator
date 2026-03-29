from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq
from dotenv import load_dotenv
import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import textwrap
import re

# load_dotenv()

# ---------------- PAGE CONFIG ----------------
st.set_page_config(page_title="GitHub README Generator", layout="wide")
st.title("🚀 GitHub README Generator")
st.caption("Generate a professional GitHub Profile README with preview and downloads.")

# ---------------- LLM ----------------
llm = ChatGroq(
    model="openai/gpt-oss-120b",
    temperature=0.2,
)

# ---------------- PROMPT ----------------
prompt = """
You are an expert GitHub profile README writer.

Your job is to generate a COMPLETE, polished, visually attractive GitHub Profile README.md.

IMPORTANT RULES:
- Return FULL README markdown only
- Do NOT explain anything outside README
- Do NOT cut output halfway
- Ensure the README is COMPLETE and properly structured
- Use only technologies relevant to the user input
- If the user gives limited input, intelligently expand it professionally
- Use markdown headings, bullets, badges, tables, emojis, and GitHub-friendly formatting
- Keep it visually attractive but not messy

README STRUCTURE:

# Header
- Greeting with user name if available
- Professional tagline
- Short intro

# About Me
Use bullet points with emojis:
- 🔭 Current work
- 🌱 Learning
- 👯 Collaboration interests
- 💬 Ask me about
- ⚡ Fun fact

# Tech Stack
Only include categories relevant to the user's input.

For each category:
- Add a heading
- Add skillicons.dev line if possible
- Add a markdown table:
| Tool | Description |

Possible categories:
- 🧠 AI / Agentic AI / GenAI
- 🔍 Vector Databases / RAG
- ⚙️ Backend / API / Realtime
- 📱 Mobile / Frontend
- 🗄 Databases / Warehouses
- 📊 Data Engineering / ETL
- 📈 Data Analysis / BI
- 🤖 Machine Learning / NLP / CV
- 🔄 AI Automation / Workflow
- 📊 MLOps / LLMOps
- 📈 Monitoring / Observability
- ☁️ DevOps / Cloud / Infra
- 🌐 Cloud Platforms / Services
- 🔐 Security / DevSecOps
- 🛡️ Guardrails / Safety

# 🚀 Current Focus
Write 6–8 bullets

# 📌 Featured Interests
Write 6–8 bullets

# 📫 Connect With Me
- GitHub link placeholder if user does not provide one
- Optional LinkedIn placeholder if appropriate

# Footer
Short inspirational footer

USER INPUT:
{readme_content}
"""

template = ChatPromptTemplate.from_messages([
    ("system", prompt)
])

# ---------------- INPUT ----------------
readme_content = st.text_area(
    "Enter your skills, tools, technologies, projects, interests, and GitHub details:",
    height=300,
    placeholder="""Example:
Name: Syed Awais Ali Shah
Role: AI Engineer, Full Stack Developer
Skills: Python, Django, FastAPI, Flutter, LangChain, LangGraph, CrewAI, Redis, PostgreSQL, Docker, Kubernetes, Terraform, Streamlit, Pandas, OpenCV, Hugging Face, GitHub Actions
Interests: AI Agents, RAG, DevOps, Mobile Apps, Automation
GitHub: https://github.com/yourusername
"""
)

# ---------------- CLEAN MARKDOWN ----------------
def clean_markdown(text):
    if not text:
        return ""
    text = text.strip()

    # Remove markdown code fences if model returns ```markdown ... ```
    text = re.sub(r"^```(?:markdown|md)?\n", "", text)
    text = re.sub(r"\n```$", "", text)

    return text.strip()

# ---------------- DOCX CREATOR ----------------
def create_readme_docx(markdown_text):
    with BytesIO() as buffer:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(40)
        section.right_margin = Pt(40)

        lines = markdown_text.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue

            # Headers
            if line.startswith("# "):
                p = doc.add_paragraph()
                run = p.add_run(line.replace("# ", ""))
                run.bold = True
                run.font.size = Pt(18)

            elif line.startswith("## "):
                p = doc.add_paragraph()
                run = p.add_run(line.replace("## ", ""))
                run.bold = True
                run.font.size = Pt(15)

            elif line.startswith("### "):
                p = doc.add_paragraph()
                run = p.add_run(line.replace("### ", ""))
                run.bold = True
                run.font.size = Pt(13)

            # Bullet points
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")

            # Table rows
            elif "|" in line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(10)

            # Normal text
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(11)

        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

# ---------------- PDF CREATOR ----------------
def create_readme_pdf(markdown_text):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    margin = 40
    y = height - 40
    line_height = 15
    wrapper = textwrap.TextWrapper(width=95)

    lines = markdown_text.split("\n")

    for line in lines:
        line_clean = line.strip()

        if not line_clean:
            y -= 8
            continue

        # Header styles
        if line_clean.startswith("# "):
            c.setFont("Helvetica-Bold", 16)
            wrapped = wrapper.wrap(line_clean.replace("# ", ""))

        elif line_clean.startswith("## "):
            c.setFont("Helvetica-Bold", 14)
            wrapped = wrapper.wrap(line_clean.replace("## ", ""))

        elif line_clean.startswith("### "):
            c.setFont("Helvetica-Bold", 12)
            wrapped = wrapper.wrap(line_clean.replace("### ", ""))

        elif line_clean.startswith("- ") or line_clean.startswith("* "):
            c.setFont("Helvetica", 11)
            wrapped = wrapper.wrap("• " + line_clean[2:])

        else:
            c.setFont("Helvetica", 10)
            wrapped = wrapper.wrap(line_clean)

        for wline in wrapped:
            c.drawString(margin, y, wline)
            y -= line_height

            if y < 50:
                c.showPage()
                y = height - 40

    c.save()
    buffer.seek(0)
    return buffer

# ---------------- GENERATE BUTTON ----------------
if st.button("✨ Generate README"):
    if not readme_content.strip():
        st.warning("Please enter your README content.")
    else:
        with st.spinner("Generating README..."):
            try:
                formatted_prompt = template.format_messages(
                    readme_content=readme_content
                )

                response = llm.invoke(formatted_prompt)
                final_output = clean_markdown(response.content)

                st.session_state.generated_readme = final_output

            except Exception as e:
                st.error(f"Error generating README: {e}")

# ---------------- SHOW RESULT ----------------
if "generated_readme" in st.session_state and st.session_state.generated_readme:
    generated = st.session_state.generated_readme

    st.success("README generated successfully!")

    tab1, tab2 = st.tabs(["👀 Preview", "📝 Raw Markdown"])

    with tab1:
        st.subheader("GitHub README Preview")
        st.markdown(generated, unsafe_allow_html=True)

    with tab2:
        st.subheader("README Markdown Code")
        st.code(generated, language="markdown")

    # ---------------- DOWNLOADS ----------------
    st.subheader("⬇️ Download README")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.download_button(
            label="Download .md",
            data=generated,
            file_name="README.md",
            mime="text/markdown"
        )

    with col2:
        st.download_button(
            label="Download .txt",
            data=generated,
            file_name="README.txt",
            mime="text/plain"
        )

    with col3:
        docx_bytes = create_readme_docx(generated)
        st.download_button(
            label="Download .docx",
            data=docx_bytes,
            file_name="README.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with col4:
        pdf_file = create_readme_pdf(generated)
        st.download_button(
            label="Download .pdf",
            data=pdf_file,
            file_name="README.pdf",
            mime="application/pdf"
        )
